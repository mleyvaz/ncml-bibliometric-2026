"""Generate the combined paper (classical + neutrosophic bibliometric
analysis of NCML, EN) as .docx in NCML journal format.
"""
from __future__ import annotations
import re
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

sys.stdout.reconfigure(encoding="utf-8")

ROOT = Path(r"C:/Users/HP/Documents/NCML_Bibliometric")
PC = ROOT / "paper_combined"
SECTIONS = PC / "sections"
OUT = PC / "NCML_Leyva_et_al_2026_Combined.docx"

# Figures can be in either of two folders
FIG_PATHS = [
    ROOT / "reports" / "clean",
    ROOT / "paper2_neutrosophic" / "figures",
]


def find_figure(name: str) -> Path | None:
    for folder in FIG_PATHS:
        p = folder / name
        if p.exists():
            return p
    return None


FONT_MAIN = "Arial"
FONT_BODY = "Times New Roman"

TITLE = ("Eight Years of Neutrosophic Computing and Machine Learning: "
         "a Bibliometric Retrospective and a Neutrosophic Extension to "
         "Bibliometric Analysis (2018–2026)")

AUTHORS = [
    ("Maikel Leyva Vazquez", 1, "*"),
    ("Yismandry Gonzalez Vargas", 2, ""),
    ("Florentin Smarandache", 3, ""),
]
AFFILIATIONS = [
    ("1", "Universidad Bolivariana del Ecuador / Universidad de Guayaquil; "
          "Editor-in-Chief, Neutrosophic Computing and Machine Learning. "
          "E-mail: mleyvaz@gmail.com"),
    ("2", "Asociacion Latinoamericana de Ciencias Neutrosoficas (ALCN), Cuba. "
          "E-mail: yismandrygonzalezvargas@gmail.com"),
    ("3", "University of New Mexico, USA; Editor-in-Chief, Neutrosophic "
          "Sets and Systems. E-mail: smarand@unm.edu"),
]


def set_font(run, name=FONT_BODY, size=10, bold=False, italic=False):
    run.font.name = name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    rFonts.set(qn("w:cs"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic


def add_para(doc, text, font=FONT_BODY, size=10, bold=False, italic=False,
             align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=4, indent_first=None):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    if indent_first is not None:
        pf.first_line_indent = Cm(indent_first)
    r = p.add_run(text)
    set_font(r, name=font, size=size, bold=bold, italic=italic)
    return p


def add_inline(paragraph, text, font=FONT_BODY, size=10, base_bold=False, base_italic=False):
    pattern = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)")
    for token in pattern.split(text):
        if not token:
            continue
        if token.startswith("**") and token.endswith("**"):
            r = paragraph.add_run(token[2:-2])
            set_font(r, name=font, size=size, bold=True, italic=base_italic)
        elif token.startswith("*") and token.endswith("*") and len(token) >= 3:
            r = paragraph.add_run(token[1:-1])
            set_font(r, name=font, size=size, bold=base_bold, italic=True)
        elif token.startswith("`") and token.endswith("`"):
            r = paragraph.add_run(token[1:-1])
            set_font(r, name="Consolas", size=size - 1, bold=base_bold, italic=base_italic)
        else:
            r = paragraph.add_run(token)
            set_font(r, name=font, size=size, bold=base_bold, italic=base_italic)


def add_h1(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(14); pf.space_after = Pt(6)
    r = p.add_run(text)
    set_font(r, name=FONT_MAIN, size=12, bold=True)


def add_h2(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(10); pf.space_after = Pt(4)
    r = p.add_run(text)
    set_font(r, name=FONT_MAIN, size=11, bold=True)


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    add_inline(p, text)


# --- Build document ---
doc = Document()
for s in doc.sections:
    s.top_margin = Cm(2.5); s.bottom_margin = Cm(2.5)
    s.left_margin = Cm(2.5); s.right_margin = Cm(2.5)

doc.styles["Normal"].font.name = FONT_BODY
doc.styles["Normal"].font.size = Pt(10)

header = doc.sections[0].header
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = hp.add_run("Neutrosophic Computing and Machine Learning, Vol. ?, ?")
set_font(r, name=FONT_MAIN, size=10, bold=True)

add_para(doc, "University of New Mexico", font=FONT_MAIN, size=10, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)

# Title
add_para(doc, TITLE, font=FONT_MAIN, size=16, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=14)

# Authors
ap = doc.add_paragraph()
ap.alignment = WD_ALIGN_PARAGRAPH.CENTER
ap.paragraph_format.space_after = Pt(4)
for i, (name, aff, mark) in enumerate(AUTHORS):
    if i > 0:
        r = ap.add_run(", "); set_font(r, name=FONT_MAIN, size=11)
    r = ap.add_run(name); set_font(r, name=FONT_MAIN, size=11)
    rs = ap.add_run(f"{aff}{mark}")
    set_font(rs, name=FONT_MAIN, size=8); rs.font.superscript = True

for aff, text in AFFILIATIONS:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(2)
    rs = p.add_run(f"{aff} "); set_font(rs, name=FONT_MAIN, size=8); rs.font.superscript = True
    r = p.add_run(text); set_font(r, name=FONT_MAIN, size=9)

add_para(doc, "*Corresponding author: mleyvaz@gmail.com",
         font=FONT_MAIN, size=9, italic=True,
         align=WD_ALIGN_PARAGRAPH.LEFT, space_after=12)


# ---- Section parser ----
HEAD_RE = re.compile(r"^(#{1,4})\s+(.+)$")
BULL_RE = re.compile(r"^\s*[-*]\s+(.+)$")
TABLE_SEP = re.compile(r"^\s*\|[\s\-:|]+\|\s*$")
TABLE_ROW = re.compile(r"^\s*\|.+\|\s*$")
QUOTE_RE = re.compile(r"^>\s+(.*)$")
IMG_RE = re.compile(r"`([^`]+\.(?:png|jpg))`")


def parse_md(doc, text, is_references=False):
    lines = text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if not line.strip():
            i += 1; continue
        m = HEAD_RE.match(line)
        if m:
            hashes, t = m.group(1), m.group(2).strip()
            (add_h1 if len(hashes) == 1 else add_h2)(doc, t)
            i += 1; continue
        if TABLE_ROW.match(line) and i + 1 < len(lines) and TABLE_SEP.match(lines[i+1]):
            header_cells = [c.strip() for c in line.strip("|").split("|")]
            i += 2
            body_rows = []
            while i < len(lines) and TABLE_ROW.match(lines[i]):
                cells = [c.strip() for c in lines[i].strip("|").split("|")]
                body_rows.append(cells); i += 1
            tbl = doc.add_table(rows=1+len(body_rows), cols=len(header_cells))
            tbl.style = "Light Grid"
            for j, c in enumerate(header_cells):
                cell = tbl.rows[0].cells[j]; cell.text = ""
                p = cell.paragraphs[0]
                r = p.add_run(c); set_font(r, name=FONT_MAIN, size=9, bold=True)
            for ri, row in enumerate(body_rows, 1):
                for j, c in enumerate(row):
                    cell = tbl.rows[ri].cells[j]; cell.text = ""
                    p = cell.paragraphs[0]
                    add_inline(p, c, size=9)
            doc.add_paragraph()
            continue
        if BULL_RE.match(line):
            txt = BULL_RE.match(line).group(1)
            if is_references:
                # References: hanging indent
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                pf = p.paragraph_format
                pf.space_after = Pt(3); pf.left_indent = Cm(0.7); pf.first_line_indent = Cm(-0.7)
                add_inline(p, txt, size=9)
            else:
                add_bullet(doc, txt)
            i += 1; continue
        # Blockquote with figure reference
        if QUOTE_RE.match(line):
            buf = []
            while i < len(lines) and QUOTE_RE.match(lines[i]):
                buf.append(QUOTE_RE.match(lines[i]).group(1)); i += 1
            txt = " ".join(buf).strip()
            for fn in IMG_RE.findall(txt):
                img = find_figure(fn)
                if img is not None:
                    pimg = doc.add_paragraph()
                    pimg.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    pimg.add_run().add_picture(str(img), width=Inches(6.3))
                else:
                    print(f"[warn] figure not found: {fn}")
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.left_indent = Cm(1); p.paragraph_format.space_after = Pt(8)
            add_inline(p, txt, size=9, base_italic=True)
            continue
        # Paragraph
        buf = [line]; i += 1
        while i < len(lines) and lines[i].strip() and not (
            HEAD_RE.match(lines[i]) or BULL_RE.match(lines[i]) or
            TABLE_ROW.match(lines[i]) or QUOTE_RE.match(lines[i])):
            buf.append(lines[i].rstrip()); i += 1
        text = " ".join(b.strip() for b in buf)
        # References entries (start with [N])
        if is_references and re.match(r"^\[\d+\]", text):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            pf = p.paragraph_format
            pf.space_after = Pt(3); pf.left_indent = Cm(0.7); pf.first_line_indent = Cm(-0.7)
            add_inline(p, text, size=9)
        else:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.first_line_indent = Cm(0.6)
            add_inline(p, text)


# Render all sections in order
files = sorted(SECTIONS.glob("*.md"))
for fname in files:
    text = fname.read_text(encoding="utf-8")
    is_refs = "reference" in fname.name.lower()
    parse_md(doc, text, is_references=is_refs)

# Footer
footer = doc.sections[0].footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = fp.add_run("M. Leyva Vazquez, Y. Gonzalez Vargas, F. Smarandache. "
               "NCML bibliometric retrospective + neutrosophic extension, 2026.")
set_font(r, name=FONT_MAIN, size=8)

doc.save(OUT)
print(f"Wrote: {OUT}")
print(f"Size: {OUT.stat().st_size / 1024:.1f} KB")
