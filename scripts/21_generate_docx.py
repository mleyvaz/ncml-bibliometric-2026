"""Generate NCML-formatted .docx from the per-section markdown drafts.

The template mimics the layout observed in NCML Spanish articles:
  - Journal header: "Neutrosophic Computing and Machine Learning, Vol. N, Year" (Arial Bold 10)
  - "University of New Mexico" line
  - Title Spanish (Arial 18)
  - Title English (Arial 18, italic)
  - Authors with superscript affiliations
  - Affiliation block with emails (Arial 9)
  - Resumen / Abstract with bold label
  - Keywords / Palabras clave (italic)
  - Numbered section headings (Arial Bold 11)
  - Body Arial 10, justified
  - References at the end

Output: paper/NCML_Leyva_et_al_2026.docx
"""
from __future__ import annotations
import re
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

sys.stdout.reconfigure(encoding="utf-8")

ROOT = Path(r"C:/Users/HP/Documents/NCML_Bibliometric")
PAPER = ROOT / "paper"
OUT = PAPER / "NCML_Leyva_et_al_2026.docx"

# ---------- Config ----------
FONT_MAIN = "Arial"
FONT_BODY = "Times New Roman"
SIZE_TITLE = 18
SIZE_SUBTITLE = 14
SIZE_HEADING1 = 12
SIZE_HEADING2 = 11
SIZE_BODY = 10
SIZE_FOOT = 9
SIZE_HEADER = 10

TITLE_ES = ("Neutrosophic Computing and Machine Learning (2018-2026): "
            "una retrospectiva bibliometrica editorial")
TITLE_EN = ("Neutrosophic Computing and Machine Learning (2018-2026): "
            "an editorial bibliometric retrospective")
VOLUME = "42"
YEAR = "2026"

AUTHORS = [
    ("Maikel Leyva Vazquez", 1, "*"),
    ("Florentin Smarandache", 2, ""),
    ("[Tercer autor pendiente — bibliometrista externo]", 3, ""),
]

AFFILIATIONS = [
    ("1", "Universidad Bolivariana del Ecuador; Editor-in-Chief, Neutrosophic "
          "Computing and Machine Learning. E-mail: mleyvaz@gmail.com"),
    ("2", "University of New Mexico, USA; Editor-in-Chief, Neutrosophic Sets "
          "and Systems. E-mail: smarand@unm.edu"),
    ("3", "[Institucion pendiente]. E-mail: [pendiente]"),
]
CORRESPONDING_LINE = "*Autor de correspondencia: mleyvaz@gmail.com"


def set_font(run, name=FONT_BODY, size=SIZE_BODY, bold=False, italic=False, color=None):
    run.font.name = name
    # Ensure East Asian font map also set (Windows quirk)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    rFonts.set(qn("w:cs"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color


def add_para(doc, text, font=FONT_BODY, size=SIZE_BODY, bold=False, italic=False,
             align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=4, space_before=0,
             first_line_indent=None):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    if first_line_indent is not None:
        pf.first_line_indent = Cm(first_line_indent)
    r = p.add_run(text)
    set_font(r, name=font, size=size, bold=bold, italic=italic)
    return p


def add_heading1(doc, text, number=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after = Pt(6)
    text_full = (f"{number}. {text}" if number else text)
    r = p.add_run(text_full)
    set_font(r, name=FONT_MAIN, size=SIZE_HEADING1, bold=True)
    return p


def add_heading2(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(8)
    pf.space_after = Pt(4)
    r = p.add_run(text)
    set_font(r, name=FONT_MAIN, size=SIZE_HEADING2, bold=True)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    pf = p.paragraph_format
    pf.space_after = Pt(2)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    # Build with mixed runs for bold inline
    _add_inline_runs(p, text, font=FONT_BODY, size=SIZE_BODY)
    return p


def _add_inline_runs(paragraph, text, font=FONT_BODY, size=SIZE_BODY, bold=False, italic=False):
    """Parse **bold** and *italic* markdown-style inline markers."""
    pattern = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)")
    for token in pattern.split(text):
        if not token:
            continue
        if token.startswith("**") and token.endswith("**"):
            r = paragraph.add_run(token[2:-2])
            set_font(r, name=font, size=size, bold=True, italic=italic)
        elif token.startswith("*") and token.endswith("*") and len(token) >= 3:
            r = paragraph.add_run(token[1:-1])
            set_font(r, name=font, size=size, bold=bold, italic=True)
        elif token.startswith("`") and token.endswith("`"):
            r = paragraph.add_run(token[1:-1])
            set_font(r, name="Consolas", size=size - 1, bold=bold, italic=italic)
        else:
            r = paragraph.add_run(token)
            set_font(r, name=font, size=size, bold=bold, italic=italic)


# ---------- Build document ----------
doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# Default style
style = doc.styles["Normal"]
style.font.name = FONT_BODY
style.font.size = Pt(SIZE_BODY)

# Header of the page
header = doc.sections[0].header
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = hp.add_run(f"Neutrosophic Computing and Machine Learning, Vol. {VOLUME}, {YEAR}")
set_font(r, name=FONT_MAIN, size=SIZE_HEADER, bold=True)

# Running header line (below header in body)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("University of New Mexico")
set_font(r, name=FONT_MAIN, size=SIZE_HEADER, bold=True)
p.paragraph_format.space_after = Pt(18)

# ---- Title (ES) ----
add_para(doc, TITLE_ES, font=FONT_MAIN, size=SIZE_TITLE, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)

# ---- Title (EN) ----
add_para(doc, TITLE_EN, font=FONT_MAIN, size=SIZE_TITLE, italic=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=16)

# ---- Authors ----
author_p = doc.add_paragraph()
author_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
author_p.paragraph_format.space_after = Pt(4)
for i, (name, aff_id, mark) in enumerate(AUTHORS):
    if i > 0:
        r = author_p.add_run(", ")
        set_font(r, name=FONT_MAIN, size=SIZE_BODY + 1)
    r = author_p.add_run(name)
    set_font(r, name=FONT_MAIN, size=SIZE_BODY + 1)
    # superscript for affiliation
    r2 = author_p.add_run(str(aff_id) + mark)
    set_font(r2, name=FONT_MAIN, size=SIZE_FOOT - 1)
    r2.font.superscript = True

# ---- Affiliations ----
for aff_id, text in AFFILIATIONS:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"{aff_id} ")
    set_font(r, name=FONT_MAIN, size=SIZE_FOOT - 1)
    r.font.superscript = True
    r = p.add_run(text)
    set_font(r, name=FONT_MAIN, size=SIZE_FOOT)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
p.paragraph_format.space_after = Pt(12)
r = p.add_run(CORRESPONDING_LINE)
set_font(r, name=FONT_MAIN, size=SIZE_FOOT, italic=True)


# ---------- Markdown parser ----------
BULLET_RE = re.compile(r"^\s*[-*]\s+(.+)$")
HEADING_RE = re.compile(r"^(#{1,4})\s+(.+)$")
TABLE_SEP_RE = re.compile(r"^\s*\|[\s\-:|]+\|\s*$")
TABLE_ROW_RE = re.compile(r"^\s*\|.+\|\s*$")
CODE_FENCE_RE = re.compile(r"^\s*```")
QUOTE_RE = re.compile(r"^>\s+(.*)$")
IMAGE_RE = re.compile(r"^!\[([^\]]*)\]\(([^)]+)\)\s*$")

FIGURES_DIR = PAPER.parent / "reports" / "clean"


def strip_tbd(text: str) -> str:
    """Remove [TBD cita X] placeholders without leaving weird spacing."""
    text = re.sub(r"\s*\[TBD cita ([^\]]+)\]", r" [\1]", text)
    return text


def parse_markdown_body(doc, md_text: str, base_section_number: int | None = None):
    """Render a section's markdown into the docx.

    We intentionally collapse frontmatter-like heading '# N. Title' to the top
    level and treat '## N.M Sub' as heading2.
    """
    lines = md_text.splitlines()
    i = 0
    in_code = False
    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip()

        # Code block
        if CODE_FENCE_RE.match(line):
            in_code = not in_code
            i += 1
            continue
        if in_code:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            r = p.add_run(line)
            set_font(r, name="Consolas", size=SIZE_BODY - 1)
            i += 1
            continue

        # Table start
        if TABLE_ROW_RE.match(line) and i + 1 < len(lines) and TABLE_SEP_RE.match(lines[i + 1]):
            # Collect table rows
            header_cells = [c.strip() for c in line.strip().strip("|").split("|")]
            i += 2  # skip separator
            body_rows = []
            while i < len(lines) and TABLE_ROW_RE.match(lines[i]):
                cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                body_rows.append(cells)
                i += 1
            # Render as docx table
            table = doc.add_table(rows=1 + len(body_rows), cols=len(header_cells))
            table.style = "Light Grid"
            for j, cell in enumerate(header_cells):
                c = table.rows[0].cells[j]
                c.text = ""
                p = c.paragraphs[0]
                r = p.add_run(cell)
                set_font(r, name=FONT_MAIN, size=SIZE_BODY - 1, bold=True)
            for ri, row in enumerate(body_rows, 1):
                for j, cell in enumerate(row):
                    c = table.rows[ri].cells[j]
                    c.text = ""
                    p = c.paragraphs[0]
                    _add_inline_runs(p, strip_tbd(cell), font=FONT_BODY, size=SIZE_BODY - 1)
            # Small space after table
            doc.add_paragraph()
            continue

        # Image: ![caption](filename.png)
        m = IMAGE_RE.match(line)
        if m:
            caption, fname = m.group(1), m.group(2).strip()
            # Resolve image: look in FIGURES_DIR
            img_path = FIGURES_DIR / fname
            if not img_path.exists():
                # Try as a straight path
                alt = Path(fname)
                if alt.exists():
                    img_path = alt
            if img_path.exists():
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(2)
                r = p.add_run()
                try:
                    r.add_picture(str(img_path), width=Inches(6.3))
                except Exception as e:
                    print(f"[warn] could not insert image {img_path}: {e}")
                if caption:
                    cap = doc.add_paragraph()
                    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cap.paragraph_format.space_after = Pt(8)
                    crun = cap.add_run(caption)
                    set_font(crun, name=FONT_MAIN, size=SIZE_FOOT, italic=True)
            else:
                print(f"[warn] figure not found: {fname}")
            i += 1
            continue

        # Heading
        m = HEADING_RE.match(line)
        if m:
            hashes, text = m.group(1), m.group(2).strip()
            text = strip_tbd(text)
            level = len(hashes)
            if level == 1:
                add_heading1(doc, text)
            elif level == 2:
                add_heading2(doc, text)
            else:
                add_heading2(doc, text)  # collapse deeper levels
            i += 1
            continue

        # Bullet
        m = BULLET_RE.match(line)
        if m:
            add_bullet(doc, strip_tbd(m.group(1)))
            i += 1
            continue

        # Blockquote (may span multiple lines). Collect full blockquote.
        if QUOTE_RE.match(line):
            quote_lines = []
            while i < len(lines) and QUOTE_RE.match(lines[i]):
                quote_lines.append(QUOTE_RE.match(lines[i]).group(1))
                i += 1
            quote_text = " ".join(quote_lines).strip()
            # Figure reference? Look for filename ending in .png/.jpg inside backticks
            figs = re.findall(r"`([^`]+\.(?:png|jpg|jpeg))`", quote_text)
            # Insert each image (centered, 6.3 inches)
            for fn in figs:
                img_path = FIGURES_DIR / fn
                if img_path.exists():
                    pimg = doc.add_paragraph()
                    pimg.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    pimg.paragraph_format.space_before = Pt(6)
                    pimg.paragraph_format.space_after = Pt(2)
                    try:
                        pimg.add_run().add_picture(str(img_path), width=Inches(6.3))
                    except Exception as e:
                        print(f"[warn] image insert failed {fn}: {e}")
                else:
                    print(f"[warn] figure missing: {fn}")
            # Then the caption
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.left_indent = Cm(1)
            p.paragraph_format.space_after = Pt(8)
            _add_inline_runs(p, strip_tbd(quote_text), font=FONT_BODY, size=SIZE_FOOT, italic=True)
            continue

        # Empty line = paragraph break
        if line.strip() == "":
            i += 1
            continue

        # Paragraph: consume until blank
        buf = [line]
        i += 1
        while i < len(lines) and lines[i].strip() and not (
            HEADING_RE.match(lines[i]) or BULLET_RE.match(lines[i]) or
            TABLE_ROW_RE.match(lines[i]) or QUOTE_RE.match(lines[i]) or
            CODE_FENCE_RE.match(lines[i])
        ):
            buf.append(lines[i].rstrip())
            i += 1
        text = strip_tbd(" ".join(b.strip() for b in buf))
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.first_line_indent = Cm(0.6)
        _add_inline_runs(p, text, font=FONT_BODY, size=SIZE_BODY)


# ---- Render each section ----
SECTION_FILES = [
    ("01_abstract.md", True),          # Abstract/keywords (special)
    ("02_introduccion.md", False),
    ("03_metodos.md", False),
    ("04_resultados.md", False),
    ("05_discusion.md", False),
    ("06_conclusion.md", False),
    ("07_recomendaciones.md", False),
]

for fname, is_abstract in SECTION_FILES:
    path = PAPER / fname
    if not path.exists():
        print(f"[warn] missing {fname}")
        continue
    text = path.read_text(encoding="utf-8")
    # Section-level title is the first # ... line inside file; we let the
    # parser handle it. For the abstract section we render as "Abstract" label
    # bolded and then the text.
    parse_markdown_body(doc, text)


# ---- References ----
add_heading1(doc, "Referencias")

# Read refs.bib and convert each entry to a numbered item
bib = (PAPER / "refs.bib").read_text(encoding="utf-8") if (PAPER / "refs.bib").exists() else ""
entries = re.findall(r"@\w+\s*\{[^@]+?\n\}", bib, re.S)


def bib_to_line(entry: str) -> str:
    key_match = re.match(r"@\w+\{([^,]+),", entry)
    key = key_match.group(1) if key_match else "?"
    fields = dict(re.findall(r"(\w+)\s*=\s*\{([^{}]*(?:\{[^}]*\}[^{}]*)*)\}", entry))
    author = fields.get("author", "").replace(" and ", ", ").replace("{", "").replace("}", "")
    title = fields.get("title", "").replace("{", "").replace("}", "")
    journal = fields.get("journal", fields.get("booktitle", fields.get("publisher", "")))
    year = fields.get("year", "")
    volume = fields.get("volume", "")
    number = fields.get("number", "")
    pages = fields.get("pages", "").replace("--", "-")
    doi = fields.get("doi", "")
    parts = [author, year, title]
    if journal:
        parts.append(journal)
    vol_str = ""
    if volume:
        vol_str = volume + ((f"({number})") if number else "")
    if vol_str:
        parts.append(vol_str)
    if pages:
        parts.append(pages)
    line = ". ".join(p for p in parts if p) + "."
    if doi:
        line += f" https://doi.org/{doi}"
    return line


for idx, entry in enumerate(entries, 1):
    line = bib_to_line(entry)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    r = p.add_run(f"[{idx}] ")
    set_font(r, name=FONT_BODY, size=SIZE_FOOT, bold=True)
    r = p.add_run(line)
    set_font(r, name=FONT_BODY, size=SIZE_FOOT)


# ---- Page footer ----
footer = doc.sections[0].footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = fp.add_run("M. Leyva Vazquez, F. Smarandache et al. Retrospectiva bibliometrica NCML 2018-2026.")
set_font(r, name=FONT_MAIN, size=SIZE_FOOT - 1)

# Save
doc.save(OUT)
print(f"Wrote: {OUT}")
print(f"Size: {OUT.stat().st_size / 1024:.1f} KB")
