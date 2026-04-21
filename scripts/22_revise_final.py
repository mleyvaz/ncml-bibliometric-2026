"""Revise the user's Final.docx:

1. Fix affiliation numbering (Smarandache: 2 -> 3).
2. Convert in-text citations to numbered [N] / [N, M] / [N-M] format.
3. Replace the GitHub placeholder with the real repo URL.
4. Remove citations that were placeholders not in the bibliography
   (Rusk & Waters 2013, McGraw 2013) and reword the surrounding sentence.
5. Add a "Disponibilidad de datos y codigo" paragraph before References.

Saves as NCML_Leyva_et_al_2026-Final-Rev.docx preserving images/tables.

Citation mapping (based on the references list in the Final):
  [1]  Lotka 1926
  [2]  Bradford 1934
  [3]  Garfield 1972
  [4]  Seglen 1992
  [5]  Haas 1992
  [6]  Newman 2005
  [7]  Tarjan 1975
  [8]  Blondel et al. 2008
  [9]  Grootendorst 2022
  [10] McInnes 2018
  [11] Reimers & Gurevych 2019
  [12] Harzing & van der Wal 2008
  [13] Martin-Martin et al. 2021
  [14] Priem et al. 2022
  [15] PRISMA 2020 (Page et al.)
  [16] Smarandache 1998
  [17] Smarandache 1999
  [18] Woodall et al. 2025
  [19] Hirsch 2005
  [20] Larsen & von Ins 2010
  [21] Petrou 2020
"""
from __future__ import annotations
import re
import sys
from pathlib import Path
from copy import deepcopy
from docx import Document

sys.stdout.reconfigure(encoding="utf-8")

ROOT = Path(r"C:/Users/HP/Documents/NCML_Bibliometric")
SRC = ROOT / "paper" / "NCML_Leyva_et_al_2026-Final.docx"
OUT = ROOT / "paper" / "NCML_Leyva_et_al_2026-Final-Rev.docx"

# GitHub repo URL (ajustable; se usa tambien en el README)
GH_REPO = "https://github.com/mleyvaz/ncml-bibliometric-2026"

# Replacements applied to every run's text (order matters: longer patterns first).
REPLACEMENTS = [
    # Placeholder URL
    (r"\[TBD URL OSF/?GitHub\]", GH_REPO),
    # Citas específicas con clave textual
    (r"\[Bradford 1934;\s*TBD cita Garfield 1972\]", "[2, 3]"),
    (r"\[Bradford 1934\s*;\s*Garfield 1972\]", "[2, 3]"),
    (r"\[Smarandache 1998,\s*1999\]", "[16, 17]"),
    (r"\[PRISMA 2020\]", "[15]"),
    (r"\[Tarjan 1975\]", "[7]"),
    (r"\[Newman 2005,\s*ec\.\s*3\.1\]", "[6]"),
    (r"\[Newman 2005\]", "[6]"),
    (r"\[Blondel et al\.\s*2008\]", "[8]"),
    (r"\[Grootendorst 2022\]", "[9]"),
    (r"\[Lotka 1926\]", "[1]"),
    (r"\[Seglen 1992\]", "[4]"),
    (r"\[Haas 1992\]", "[5]"),
    (r"\[Petrou 2020\]", "[21]"),
    (r"\[Woodall et al\.\s*\(?2025\)?\]", "[18]"),
    (r"Woodall et al\.\s*\(2025\)", "Woodall et al. [18]"),
    # TBD residual
    (r"\[TBD cita\s*\]", ""),
    (r"\[TBD cita ([^\]]+)\]", r""),
    # Citas de año solo (Harzing y Martin-Martin en contexto "apellido [año]")
    (r"Harzing y van der Wal\s*\[2008\]", "Harzing y van der Wal [12]"),
    (r"Martin-Martin et al\.\s*\[2020\]", "Martin-Martin et al. [13]"),
    # Remover citas espúreas (Rusk y McGraw) y reformular
    (r"\s*—?\s*la psicologia positiva\s*\[Rusk\s*&\s*Waters 2013\][^—]*\[McGraw 2013\][^—]*—\s*",
     " — niveles similares aparecen documentados en otros nichos especializados jovenes — "),
    (r"\[Rusk\s*&\s*Waters 2013\]", ""),
    (r"\[McGraw 2013\]", ""),
    # Doble espacios generados por eliminaciones
    (r" +\.", "."),
    (r" +,", ","),
    (r"  +", " "),
]


def apply_replacements(text: str) -> str:
    for pat, rep in REPLACEMENTS:
        text = re.sub(pat, rep, text)
    return text


doc = Document(SRC)


# ---- 1. Affiliation numbering fix ----
# The user wrote "Florentin Smarandache2" but his affiliation is at index 3.
# We rewrite paragraph 3 (authors line) so Smarandache carries the superscript 3.
# We also rewrite paragraph 5 so ALCN affiliation is clearly labelled "2".
def fix_authors_line(p):
    # Expected text approximately:
    # "Maikel Leyva Vazquez1*, Yismandry González Vargas2, Florentin Smarandache2"
    current = p.text
    new = current.replace("Florentin Smarandache2", "Florentin Smarandache3")
    new = new.replace("Florentin Smarandache 2", "Florentin Smarandache 3")
    if new != current:
        # Wipe runs and rewrite with one run preserving initial font
        original_run = p.runs[0] if p.runs else None
        font_name = original_run.font.name if original_run else "Arial"
        font_size = original_run.font.size if original_run else None
        # Clear all runs
        for r in list(p.runs):
            r._element.getparent().remove(r._element)
        r = p.add_run(new)
        if font_name:
            r.font.name = font_name
        if font_size:
            r.font.size = font_size
        r.font.bold = False


for p in doc.paragraphs[:8]:
    if "Florentin Smarandache" in p.text and "2" in p.text and "Maikel" in p.text:
        fix_authors_line(p)


# ---- 2. Replacements in text (paragraphs + tables) ----
def rewrite_paragraph(p):
    """Apply text replacements while trying to preserve per-run formatting.
    When the replacement target is inside a single run, we modify run.text.
    When it spans multiple runs, we concatenate, replace, and rewrite the first
    run with the new text (clearing the rest). Citation markers don't need
    inline bold/italic, so this is acceptable.
    """
    original = p.text
    new_text = apply_replacements(original)
    if new_text == original:
        return False
    # Try per-run replacement first (simple case)
    per_run_ok = True
    test_runs = [apply_replacements(r.text) for r in p.runs]
    if "".join(test_runs) == new_text:
        # Safe per-run replacement
        for r, nt in zip(p.runs, test_runs):
            r.text = nt
        return True
    # Fallback: rewrite paragraph preserving first run's font
    first = p.runs[0] if p.runs else None
    font_name = first.font.name if first else None
    font_size = first.font.size if first else None
    font_bold = first.font.bold if first else None
    font_italic = first.font.italic if first else None
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    run = p.add_run(new_text)
    if font_name:
        run.font.name = font_name
    if font_size:
        run.font.size = font_size
    if font_bold is not None:
        run.font.bold = font_bold
    if font_italic is not None:
        run.font.italic = font_italic
    return True


changed = 0
for p in doc.paragraphs:
    if rewrite_paragraph(p):
        changed += 1

for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                if rewrite_paragraph(p):
                    changed += 1

print(f"Paragraphs rewritten: {changed}")


# ---- 3. Add Data Availability paragraph before References ----
# Find the paragraph that equals "Referencias"
insert_idx = None
for i, p in enumerate(doc.paragraphs):
    if p.text.strip().lower() == "referencias":
        insert_idx = i
        break

if insert_idx is not None:
    # Build the new paragraphs to insert
    body = doc.paragraphs[insert_idx]._element

    # Create the section heading "Disponibilidad de datos y codigo"
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt

    def make_paragraph(text, bold=False, size=None, font="Arial"):
        p = OxmlElement("w:p")
        pPr = OxmlElement("w:pPr")
        p.append(pPr)
        r = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        rFonts = OxmlElement("w:rFonts")
        rFonts.set(qn("w:ascii"), font)
        rFonts.set(qn("w:hAnsi"), font)
        rPr.append(rFonts)
        if bold:
            b = OxmlElement("w:b")
            rPr.append(b)
        if size is not None:
            sz = OxmlElement("w:sz")
            sz.set(qn("w:val"), str(size * 2))
            rPr.append(sz)
        r.append(rPr)
        t = OxmlElement("w:t")
        t.set(qn("xml:space"), "preserve")
        t.text = text
        r.append(t)
        p.append(r)
        return p

    heading = make_paragraph("Disponibilidad de datos y codigo", bold=True, size=12, font="Arial")
    body.addprevious(heading)

    body_text = (
        "El pipeline completo (scraping, descarga de PDF, extraccion de texto, "
        "desambiguacion, ajustes de Lotka/Bradford, red de coautoria, modelado "
        "de topicos) esta disponible como codigo abierto en "
        f"{GH_REPO} bajo licencias MIT (codigo) y CC-BY 4.0 (datos). El "
        "repositorio contiene: los 22 scripts numerados reproducibles; el "
        "archivo Excel NCML_bibliometric_dataset.xlsx con 19 hojas que "
        "consolidan todas las tablas analiticas; los backups ZIP con los 728 "
        "PDFs descargados y con los datos intermedios; y las figuras del "
        "estudio en formato PNG 170 dpi. Los investigadores interesados en "
        "verificar, replicar o extender el analisis pueden clonar el "
        "repositorio y ejecutar los scripts en orden numerico (01 a 22). "
        "Los datos intermedios se publican con hashes SHA-256 para verificacion "
        "de integridad."
    )
    para = make_paragraph(body_text, size=10, font="Times New Roman")
    body.addprevious(para)
    print("Inserted 'Disponibilidad de datos y codigo' section before References.")
else:
    print("[warn] Could not locate 'Referencias' paragraph for insertion.")


# ---- 4. Save ----
doc.save(OUT)
print(f"\nSaved: {OUT}")
print(f"Size: {OUT.stat().st_size / 1024:.1f} KB")
