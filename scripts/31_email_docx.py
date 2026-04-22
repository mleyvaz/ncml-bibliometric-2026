"""Generate the Smarandache email as .docx for copy-paste convenience."""
import sys, re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

sys.stdout.reconfigure(encoding="utf-8")

ROOT = Path(r"C:/Users/HP/Documents/NCML_Bibliometric/paper_combined")
MD = ROOT / "email_to_smarandache.md"
OUT = ROOT / "email_to_smarandache.docx"

doc = Document()
for s in doc.sections:
    s.top_margin = Cm(2); s.bottom_margin = Cm(2)
    s.left_margin = Cm(2.5); s.right_margin = Cm(2.5)
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)


def add(text, bold=False, italic=False, size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
        space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    # inline bold/italic via ** and *
    parts = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*)", text)
    for t in parts:
        if not t:
            continue
        if t.startswith("**") and t.endswith("**"):
            r = p.add_run(t[2:-2])
            r.font.bold = True
        elif t.startswith("*") and t.endswith("*") and len(t) >= 3:
            r = p.add_run(t[1:-1])
            r.font.italic = True
        else:
            r = p.add_run(t)
            r.font.bold = bold; r.font.italic = italic
        r.font.size = Pt(size); r.font.name = "Calibri"
    return p


text = MD.read_text(encoding="utf-8")
lines = text.splitlines()
i = 0
while i < len(lines):
    line = lines[i].rstrip()
    if line.startswith("# "):
        add(line[2:].strip(), bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=14)
    elif line.startswith("## "):
        add(line[3:].strip(), bold=True, size=12, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=6)
    elif line.startswith("---"):
        add("—" * 40, align=WD_ALIGN_PARAGRAPH.CENTER)
    elif line.startswith("- "):
        p = doc.add_paragraph(style="List Bullet")
        parts = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*)", line[2:].strip())
        for t in parts:
            if not t:
                continue
            if t.startswith("**") and t.endswith("**"):
                r = p.add_run(t[2:-2]); r.font.bold = True
            elif t.startswith("*") and t.endswith("*") and len(t) >= 3:
                r = p.add_run(t[1:-1]); r.font.italic = True
            else:
                r = p.add_run(t)
            r.font.name = "Calibri"; r.font.size = Pt(11)
    elif re.match(r"^\d+\.\s", line):
        p = doc.add_paragraph(style="List Number")
        parts = re.split(r"(\*\*[^*]+\*\*)", re.sub(r"^\d+\.\s+", "", line))
        for t in parts:
            if not t: continue
            if t.startswith("**") and t.endswith("**"):
                r = p.add_run(t[2:-2]); r.font.bold = True
            else:
                r = p.add_run(t)
            r.font.name = "Calibri"; r.font.size = Pt(11)
    elif not line.strip():
        pass
    else:
        # paragraph: collect
        buf = [line]; i += 1
        while i < len(lines) and lines[i].strip() and not (
            lines[i].startswith("#") or lines[i].startswith("- ") or
            lines[i].startswith("---") or re.match(r"^\d+\.\s", lines[i])):
            buf.append(lines[i].rstrip()); i += 1
        add(" ".join(b.strip() for b in buf))
        continue
    i += 1

doc.save(OUT)
print(f"Wrote: {OUT}")
